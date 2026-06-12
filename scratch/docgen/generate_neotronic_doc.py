from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "NEOTRONIC_Project_Documentation.docx"


BLUE = RGBColor(31, 78, 121)
MID_BLUE = RGBColor(46, 116, 181)
LIGHT_FILL = "E8EEF5"
GRAY_FILL = "F2F4F7"


modules = [
    ("Authentication & User Management", "Login, logout, JWT issue/validation, route guards, users, roles, menu selection, and password lifecycle controls."),
    ("Dashboard", "Role-specific KPIs, operational summaries, widgets, sales indicators, stock indicators, and executive analytics."),
    ("Lead Management", "Lead capture, assignment, follow-up, meetings, activity log, quote tracking, lost-lead analysis, and lead-stage workflow."),
    ("Customer Management", "Customer/client account registration, profiles, account history, contact details, and relationship tracking."),
    ("Sales Management", "Quotations, sales orders, invoices, delivery orders, sales returns, receipt vouchers, debit/credit notes, and payment references."),
    ("Purchase Management", "Supplier/vendor management, purchase orders, goods receipt, purchase master/details, purchase returns, and GRN VAT reporting."),
    ("Inventory Management", "Item master, brand, model, HSN, stock, stock add, stock details, stock transfer/adjustment, stock take, and stock reports."),
    ("Requirement Management", "Requirement master/details, lead requirements, workflow status, price requests, price responses, and approval tracking."),
    ("Project Management", "Operational task allocation, department/status tracking, resource ownership, progress monitoring, and management reporting."),
    ("Reports Module", "Sales, purchase, stock, ledger, outstanding, journal, VAT, GRN VAT, statement of account, profit and loss, and export workflows."),
    ("Settings Module", "Company, financial year, account groups, terms, currency, country, designation, department, vertical, user type, and general settings."),
]

tables = [
    ("User_Details", "Stores employee/user login identity and profile information.", "User_Details_Id", "User_Role, User_Type, Department"),
    ("User_Role", "Defines role names and role-level access grouping.", "User_Role_Id", "User_Details, User_Menu_Selection"),
    ("User_Menu_Selection", "Maps users/roles to permitted menus and screens.", "User_Menu_Selection_Id", "Menu, User_Details"),
    ("Menu", "Stores navigation/menu records consumed by the Angular admin shell.", "Menu_Id", "User_Menu_Selection"),
    ("Lead", "Stores CRM lead records, lifecycle status, assignment, follow-up, lost-lead fields, and contact metadata.", "Lead_Id", "Status, User_Details, Department, Vertical"),
    ("Follow_up", "Stores lead follow-up records and next action tracking.", "FollowUp_Id", "Lead, User_Details"),
    ("lead_activity_log", "Stores audit-style CRM activity records for lead creation, remarks, stage changes, and follow-ups.", "LeadActivityLog_Id", "Lead, User_Details"),
    ("lead_meeting", "Stores online/offline lead meeting notes and outcomes.", "LeadMeeting_Id", "Lead, User_Details"),
    ("lead_quote_tracking", "Stores quote-sent tracking for CRM leads and requirements.", "LeadQuoteTracking_Id", "Lead, requirementmaster"),
    ("Client_Accounts", "Stores customer/client account records.", "Client_Accounts_Id", "Sales, Receipts, Ledger"),
    ("Company", "Stores company master data and organizational settings.", "Company_Id", "Company-wide transactions"),
    ("Accounts", "Stores chart of accounts and ledger accounts.", "Accounts_Id", "Account_Group, transactions"),
    ("Account_Group", "Groups accounts for reporting and accounting controls.", "Account_Group_Id", "Accounts"),
    ("Account_Years", "Stores financial/accounting year definitions.", "Account_Years_Id", "Financial transactions"),
    ("Item", "Stores product/item master data.", "Item_Id", "Item_Group, Brand, Model, HSN, Stock"),
    ("Item_Group", "Groups product items.", "Item_Group_Id", "Item"),
    ("Brand", "Stores item brand master data.", "Brand_Id", "Item"),
    ("Model", "Stores product model master data.", "Model_Id", "Item"),
    ("HSN", "Stores HSN/tax classification codes.", "HSN_Id", "Item, Tax"),
    ("Stock", "Stores stock master/balance records.", "Stock_Id", "Item, Stock_Details"),
    ("Stock_Details", "Stores stock transaction details.", "Stock_Details_Id", "Stock, Item"),
    ("Stock_Add_Master", "Stores stock addition header records.", "Stock_Add_Master_Id", "Stock_Add_Details"),
    ("Stock_Add_Details", "Stores stock addition line items.", "Stock_Add_Details_Id", "Stock_Add_Master, Item"),
    ("Stock_Take_Master", "Stores physical stock-take headers.", "Stock_Take_Master_Id", "stock_take_name, Stock_Details"),
    ("stock_take_name", "Stores stock-take session names.", "stock_take_name_Id", "Stock_Take_Master"),
    ("requirementmaster", "Stores requirement header records.", "Requirement_Id", "requirementdetails, LeadRequirement"),
    ("requirementdetails", "Stores requirement item/detail records.", "Requirement_Details_Id", "requirementmaster, Item"),
    ("requirementworkflow", "Stores requirement approval/status history.", "RequirementWorkflow_Id", "requirementmaster, User_Details"),
    ("Price_Response", "Stores vendor/internal price responses against requirements.", "Price_Response_Id", "requirementmaster, Item"),
    ("Sales_Master", "Stores sales transaction headers.", "Sales_Master_Id", "Sales_Details, Client_Accounts"),
    ("Sales_Details", "Stores sales transaction line items.", "Sales_Details_Id", "Sales_Master, Item"),
    ("salesquotationmaster", "Stores quotation headers.", "SalesQuotationMaster_Id", "salesquotationdetails, Client_Accounts"),
    ("salesquotationdetails", "Stores quotation line items.", "SalesQuotationDetails_Id", "salesquotationmaster, Item"),
    ("salesordermaster", "Stores sales order headers.", "SalesOrderMaster_Id", "salesorderdetails, Client_Accounts"),
    ("salesorderdetails", "Stores sales order line items.", "SalesOrderDetails_Id", "salesordermaster, Item"),
    ("deliveryordermaster", "Stores delivery order headers.", "DeliveryOrderMaster_Id", "deliveryorderdetails, Sales"),
    ("deliveryorderdetails", "Stores delivery order details.", "DeliveryOrderDetails_Id", "deliveryordermaster, Item"),
    ("Purchase_Master", "Stores purchase transaction headers.", "Purchase_Master_Id", "Purchase_Details, supplier account"),
    ("Purchase_Details", "Stores purchase transaction line items.", "Purchase_Details_Id", "Purchase_Master, Item"),
    ("Purchase_Order_Master", "Stores purchase order headers.", "Purchase_Order_Master_Id", "purchaseorderdetails"),
    ("purchaseorderdetails", "Stores purchase order line items.", "PurchaseOrderDetails_Id", "Purchase_Order_Master, Item"),
    ("Receipt_Voucher", "Stores customer receipt vouchers.", "Receipt_Voucher_Id", "Receipt_Reference, Accounts"),
    ("Payment_Voucher", "Stores supplier/payment vouchers.", "Payment_Voucher_Id", "Payment_Reference, Accounts"),
    ("Journal_Entry", "Stores journal accounting entries.", "Journal_Entry_Id", "Accounts"),
    ("Contra_Entry", "Stores cash/bank contra transfers.", "Contra_Entry_Id", "Accounts"),
]

api_groups = [
    ("POST", "/Login/Login_Checks?userName={user}&password={password}", "Authenticate user and return JWT token with user result set.", "Public"),
    ("POST", "/Lead/Save_Lead/", "Create or update lead, assignment, follow-up, lost-lead fields, and activity log.", "Bearer JWT"),
    ("GET", "/Lead/Get_Leads/", "Return lead list, enriched with latest follow-up details.", "Bearer JWT"),
    ("GET", "/Lead/Get_Dropdowns_Lead/", "Return lead dropdown data including seeded stages, verticals, and designations.", "Bearer JWT"),
    ("GET", "/Lead/Delete_Lead/{Lead_Id}", "Soft-delete lead where DeleteStatus exists; otherwise delete record.", "Bearer JWT"),
    ("GET", "/Lead/Get_Lead_FollowUp_History/{Lead_Id}", "Return follow-up history for a lead.", "Bearer JWT"),
    ("GET", "/Lead/Get_Lead_Activity_Log/{Lead_Id}", "Return CRM activity log for a lead.", "Bearer JWT"),
    ("POST", "/Lead/Save_Lead_Meeting/", "Add online/offline meeting note and log activity.", "Bearer JWT"),
    ("POST", "/Lead/Save_Lead_Quote_Tracking/", "Add quote tracking information against lead/requirement.", "Bearer JWT"),
    ("POST", "/requirementmaster/Save_*", "Create/update requirement master records.", "Bearer JWT"),
    ("POST", "/requirementdetails/Save_*", "Create/update requirement line items.", "Bearer JWT"),
    ("POST", "/Price_Response/Save_*", "Capture pricing response records.", "Bearer JWT"),
    ("POST", "/salesquotationmaster/Save_*", "Create/update quotation header.", "Bearer JWT"),
    ("POST", "/salesquotationdetails/Save_*", "Create/update quotation details.", "Bearer JWT"),
    ("POST", "/salesordermaster/Save_*", "Create/update sales order header.", "Bearer JWT"),
    ("POST", "/salesorderdetails/Save_*", "Create/update sales order details.", "Bearer JWT"),
    ("POST", "/Sales_Master/Save_*", "Create/update sales invoice/master transaction.", "Bearer JWT"),
    ("GET", "/Sales_Master/Print_Quotation/{id}", "Public print route allowed by JWT middleware exception.", "Public/controlled"),
    ("POST", "/Purchase_Order_Master/Save_*", "Create/update purchase order header.", "Bearer JWT"),
    ("POST", "/purchaseorderdetails/Save_*", "Create/update purchase order details.", "Bearer JWT"),
    ("POST", "/Stock/Save_*", "Create/update stock records.", "Bearer JWT"),
    ("POST", "/Stock_Take_Master/Save_*", "Create/update stock-take transaction.", "Bearer JWT"),
    ("GET", "/Dashboard/*", "Return KPI and dashboard data sets.", "Bearer JWT"),
    ("POST", "/User_Details/Save_*", "Create/update user profile.", "Bearer JWT"),
    ("POST", "/User_Role/Save_*", "Create/update role records.", "Bearer JWT"),
    ("POST", "/User_Menu_Selection/Save_*", "Create/update user/menu permissions.", "Bearer JWT"),
]


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(9)
    if color:
        r.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_fill(hdr[i], LIGHT_FILL)
        set_cell_text(hdr[i], h, bold=True, color=BLUE)
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_h3(doc, text):
    doc.add_heading(text, level=3)


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_fill(cell, "F8FAFC")
    cell.text = ""
    p = cell.paragraphs[0]
    for idx, line in enumerate(lines):
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
        if idx != len(lines) - 1:
            r.add_break()
    doc.add_paragraph()


def configure(doc):
    sec = doc.sections[0]
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color in [("Heading 1", 16, MID_BLUE), ("Heading 2", 13, MID_BLUE), ("Heading 3", 11.5, BLUE)]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(10 if name != "Heading 1" else 14)
        st.paragraph_format.space_after = Pt(5)


def page_break(doc):
    doc.add_page_break()


def cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NEOTRONIC")
    r.bold = True
    r.font.size = Pt(32)
    r.font.color.rgb = BLUE
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Enterprise Resource Planning System")
    r.font.size = Pt(17)
    r.font.color.rgb = MID_BLUE
    doc.add_paragraph()
    add_table(doc, ["Document Attribute", "Value"], [
        ("Project Name", "NEOTRONIC"),
        ("Company Name", "NEOTRONIC / Client Organization"),
        ("Document Type", "Software Project Documentation and Client Handover Manual"),
        ("Version", "1.0"),
        ("Prepared By", "Project Documentation Team"),
        ("Prepared Date", date.today().strftime("%d %B %Y")),
        ("Confidentiality", "Confidential - intended for authorized client, management, audit, investor, and development stakeholders only."),
    ], [2.0, 4.2])
    add_para(doc, "This document contains technical, operational, architectural, and business information related to the NEOTRONIC ERP system. Distribution, copying, and disclosure must be controlled by the owning organization.")
    page_break(doc)


def intro_sections(doc):
    add_h1(doc, "1. Document Control")
    add_table(doc, ["Version", "Date", "Author", "Description", "Status"], [
        ("0.1", "Draft", "Documentation Team", "Initial outline and section planning.", "Internal"),
        ("0.9", "Review", "Project Team", "Client-handover sections completed.", "Review"),
        ("1.0", date.today().strftime("%d %b %Y"), "Project Documentation Team", "First formal issue for stakeholder submission.", "Released"),
    ], [0.8, 1.0, 1.4, 2.5, 0.8])
    add_table(doc, ["Role", "Name/Department", "Responsibility", "Approval"], [
        ("Project Sponsor", "Client Management", "Business ownership and budget approval.", "Pending signature"),
        ("Product Owner", "Operations/ERP Lead", "Functional acceptance and prioritization.", "Pending signature"),
        ("Technical Lead", "Development Team", "Architecture and implementation acceptance.", "Pending signature"),
        ("QA Lead", "Testing Team", "Verification and release quality acceptance.", "Pending signature"),
        ("Security/Audit", "Compliance Team", "Security, auditability, and control review.", "Pending signature"),
    ], [1.25, 1.45, 2.5, 1.1])
    add_h1(doc, "2. Executive Summary")
    add_para(doc, "NEOTRONIC is an enterprise web-based ERP system designed to centralize customer relationship management, lead tracking, sales, purchases, inventory, accounting references, reporting, and administrative configuration. The system uses an Angular frontend, a Node.js/Express backend, a MySQL database, REST-style APIs, and JWT-protected role-based access control.")
    add_bullets(doc, [
        "Improve operational visibility across sales, purchase, stock, requirements, accounts, and management reporting.",
        "Reduce manual data duplication by maintaining shared master data for customers, users, items, accounts, departments, roles, and settings.",
        "Provide a controlled API layer with authentication, authorization, error handling, and database transaction boundaries.",
        "Support future growth through modular Angular screens, route-based backend resources, stored procedures, and database indexing.",
    ])
    add_h1(doc, "3. Project Overview")
    add_para(doc, "The project is positioned as a client-facing ERP platform for daily business execution. It combines operational workflow screens with master data maintenance and reporting. The repository contains a frontend Angular application under `frontend`, a backend Express application under `backend`, database migration/procedure scripts, upload/template utilities, and verification scripts used during development.")
    add_table(doc, ["Area", "Scope"], [
        ("In Scope", "ERP web application, REST APIs, authentication, role/menu control, CRM leads, customers, sales, purchases, inventory, requirements, reports, settings, deployment and maintenance guidance."),
        ("Out of Scope", "Native mobile store release, third-party payment gateway settlement, full accounting statutory filing, production infrastructure procurement, and data migration execution unless separately contracted."),
        ("Primary Users", "Super Admin, Admin, Manager, Sales Executive, Inventory Manager, Accounts Team, Employee."),
        ("Business Value", "Centralized transactions, reduced handoff friction, controlled approvals, better audit traceability, and faster access to operational data."),
    ], [1.3, 5.1])
    page_break(doc)


def architecture(doc):
    add_h1(doc, "4. System Architecture")
    add_para(doc, "NEOTRONIC follows a layered web architecture. Angular handles presentation, navigation, forms, and service calls. Node.js/Express exposes route-based APIs and middleware. MySQL stores master and transaction data, with many write/read operations delegated to stored procedures.")
    add_code_block(doc, [
        "+--------------------+       HTTPS/REST        +------------------------+",
        "| Angular Frontend   |  ---------------------> | Node.js Express API    |",
        "| Components/Routes  |                         | Routes/Models/Middleware|",
        "+---------+----------+                         +-----------+------------+",
        "          |                                                |",
        "          | JWT Bearer Token                               | mysql2 pool",
        "          v                                                v",
        "+--------------------+                         +------------------------+",
        "| Browser Session    |                         | MySQL neotronics_db    |",
        "| Local app state    |                         | Tables + Procedures    |",
        "+--------------------+                         +------------------------+",
    ])
    add_h2(doc, "Frontend Architecture")
    add_para(doc, "The frontend is an Angular 8 application using routed modules, services, guards, Angular Material/Bootstrap UI components, RxJS, charting/export libraries, and Cordova packaging artifacts. The root router redirects empty traffic to `auth/login`, lazy-loads `AuthModule`, and lazy-loads the guarded `AdminModule` for authenticated users.")
    add_bullets(doc, [
        "Components: each business screen is represented by a module component under `frontend/src/app/modules/admin` or `frontend/src/app/modules/auth`.",
        "Services: API wrappers under `frontend/src/app/services` centralize HTTP calls for routes such as Lead, Stock, Sales, Purchase, User, and Settings resources.",
        "Routing: `app.routing.ts` applies lazy loading and `useHash: true`; admin routing is protected with `CanAdminGuard`.",
        "Guards: `CanAdminGuard` and auth guards protect route access based on authentication/session state.",
        "State management: primarily component/service state using RxJS observables and browser/session storage; a formal NgRx store is not present in the inspected package.",
        "Lazy loading: auth and admin modules are loaded by route, reducing the initial shell surface and keeping authentication separated from ERP screens.",
    ])
    add_h2(doc, "Backend Architecture")
    add_para(doc, "The backend is an Express application listening through `backend/bin/ubillmlm` with `app.js` registering middleware, public login routes, JWT middleware, and module route mounts. The data access layer is organized as paired `routes` and `models` files. Routes validate/shape HTTP input; models execute MySQL queries and stored procedures through a mysql2 pool.")
    add_bullets(doc, [
        "Controllers/routes: Express routers such as `Lead.js`, `Sales_Master.js`, `Purchase_Master.js`, `Stock.js`, `User_Details.js`, and `Dashboard.js` expose endpoint groups.",
        "Models/services: model files contain database calls, stored procedure calls, helper seed logic, and transaction logic for module operations.",
        "Middleware: CORS, body parser, cookie parser, request context, auto response wrapper, Morgan logging, JWT validation, and central error handling.",
        "Authentication: login route signs JWT tokens with HS256 using the configured secret; protected routes require a Bearer token.",
        "Authorization: role/menu access is implemented through user roles, user types, menu selection, Angular guards, and backend protected routes.",
        "API flow: request -> Express middleware -> route -> model/stored procedure -> MySQL -> response wrapper/error handler -> Angular service.",
    ])
    add_h2(doc, "Database Architecture")
    add_para(doc, "The MySQL database `neotronics_db` contains master tables, transaction headers, transaction details, workflow tables, lookup tables, and stored procedures. Most business objects use an integer primary key, a soft-delete convention through `DeleteStatus` where available, and foreign-key-style relationships by identifier columns.")
    add_bullets(doc, [
        "Master data: users, roles, menus, company, accounts, items, brands, models, departments, designations, country, HSN, currency, terms, payment terms.",
        "Transactions: sales, purchase, stock, journal, contra, receipt, payment, quotation, order, delivery, returns, requirements, price responses.",
        "Relationships: header/detail relationships are common, for example Sales_Master to Sales_Details and Purchase_Order_Master to purchaseorderdetails.",
        "Constraints: enforce required fields at API/UI level and database level; use primary keys, indexes, and stored procedure checks for business constraints.",
        "Indexing: prioritize lookup/search fields, foreign-key columns, status fields, date fields, and soft-delete flags.",
    ])
    page_break(doc)


def tech_stack(doc):
    add_h1(doc, "5. Technology Stack")
    add_table(doc, ["Layer", "Technology", "Version/Package", "Purpose"], [
        ("Frontend", "Angular", "8.0.0", "Single-page ERP web application."),
        ("Frontend", "Angular Router", "8.0.0", "Lazy module routing and guarded navigation."),
        ("Frontend", "Angular Material/CDK", "8.x", "UI controls, dialogs, forms, and layouts."),
        ("Frontend", "Bootstrap / Material Dashboard", "4.x", "Responsive admin styling and layout system."),
        ("Frontend", "RxJS", "6.5.2", "Async HTTP and UI state handling."),
        ("Frontend", "xlsx / ts-xlsx / file-saver", "0.18.x / 0.0.11 / 2.0.5", "Spreadsheet import/export and downloads."),
        ("Backend", "Node.js", "Runtime", "Server-side JavaScript runtime."),
        ("Backend", "Express", "4.17.1", "REST API routing and middleware pipeline."),
        ("Backend", "mysql2/mysql", "3.22.2 / 2.11.1", "MySQL connection pool and query execution."),
        ("Backend", "express-jwt/jsonwebtoken", "8.5.1 / direct dependency", "JWT validation and token signing."),
        ("Backend", "multer", "2.1.1", "File upload handling."),
        ("Backend", "morgan", "1.9.1", "HTTP request logging."),
        ("Backend", "puppeteer/angular-pdf-generator", "24.43.0 / 0.1.0", "PDF/document generation support."),
        ("Database", "MySQL", "Configured database: neotronics_db", "Relational master, transaction, workflow, and reporting data store."),
        ("Deployment", "Nginx", "Recommended", "TLS termination, static Angular hosting, reverse proxy to Node API."),
        ("Deployment", "PM2/systemd", "Recommended", "Node.js process management."),
    ], [1.0, 1.35, 1.65, 2.5])
    page_break(doc)


def functional_modules(doc):
    add_h1(doc, "6. Functional Modules")
    for idx, (name, summary) in enumerate(modules, start=1):
        add_h2(doc, f"6.{idx} {name}")
        add_para(doc, summary)
        add_table(doc, ["Capability", "Description", "Controls / Outputs"], [
            ("Create / Capture", f"Users can create core {name.lower()} records through Angular forms and backend route calls.", "Required field validation, duplicate checks where applicable, success/error response."),
            ("Edit / Update", "Authorized users can revise operational details while preserving the relevant identifiers and workflow context.", "Audit fields, update stored procedure, status refresh."),
            ("Delete / Close", "Records use soft-delete or controlled cancellation where business history must remain available.", "DeleteStatus or cancellation marker; permission check."),
            ("Search / Track", "Users can search, filter, and track status, owners, dates, references, and remarks.", "Grid/list views, status labels, dashboard/report visibility."),
            ("Reporting", "Module data contributes to role-level dashboards, operational reports, or export files.", "PDF/Excel export, report filters, KPI feeds."),
        ], [1.2, 3.2, 1.9])
        add_bullets(doc, [
            f"Primary users: {', '.join(['Admin', 'Manager'])}; additional access depends on role/menu permissions.",
            "Data quality expectations: required identifiers, consistent date formats, numeric validation, and controlled status transitions.",
            "Audit expectations: created/updated metadata and activity history should be retained for important workflow actions.",
        ])
        add_h3(doc, "Operational Workflow")
        add_numbered(doc, [
            f"Open the {name} screen from the authorized admin menu.",
            "Create or search the relevant record using business identifiers and filters.",
            "Validate mandatory fields, related master data, numeric values, dates, and workflow status.",
            "Save the record through the corresponding Angular service and Express route.",
            "Review confirmation, refresh the listing, and verify report/dashboard impact where applicable.",
        ])
        add_h3(doc, "Acceptance Criteria")
        add_bullets(doc, [
            "Only authorized users can access create, edit, delete, approval, and reporting actions.",
            "Required fields are validated before submission and repeated on the API/database side where practical.",
            "Successful save operations return a stable identifier and a user-safe success message.",
            "Failed operations return a user-safe error message without exposing database credentials or internal stack traces.",
            "Reports and exports reconcile with saved transaction and master data.",
        ])
        page_break(doc)


def database_doc(doc):
    add_h1(doc, "7. Database Documentation")
    add_para(doc, "The following table documents representative and inspected database entities. Exact physical columns should be confirmed from the production schema dump before final database sign-off; this document records the expected purpose, keys, and relationships based on repository models, routes, SQL migration files, and naming conventions.")
    add_table(doc, ["Table Name", "Purpose", "Primary Key", "Relationships"], tables, [1.55, 2.7, 1.15, 1.1])
    add_h2(doc, "ER Diagram Explanation")
    add_code_block(doc, [
        "User_Role 1---* User_Details 1---* User_Menu_Selection *---1 Menu",
        "Lead *---1 User_Details       Lead 1---* Follow_up",
        "Lead 1---* lead_activity_log  Lead 1---* lead_meeting",
        "Lead 1---* lead_quote_tracking *---1 requirementmaster",
        "Client_Accounts 1---* Sales_Master 1---* Sales_Details *---1 Item",
        "Client_Accounts 1---* salesquotationmaster 1---* salesquotationdetails *---1 Item",
        "Purchase_Master 1---* Purchase_Details *---1 Item",
        "Purchase_Order_Master 1---* purchaseorderdetails *---1 Item",
        "Item_Group 1---* Item *---1 Brand",
        "Item *---1 Model      Item *---1 HSN",
        "Item 1---* Stock_Details *---1 Stock",
        "requirementmaster 1---* requirementdetails *---1 Item",
        "requirementmaster 1---* requirementworkflow",
    ])
    add_h2(doc, "Indexing Strategy")
    add_bullets(doc, [
        "Create indexes on all header/detail join columns such as Sales_Master_Id, Purchase_Master_Id, Requirement_Id, Lead_Id, and Item_Id.",
        "Index common search columns such as names, document numbers, dates, status identifiers, and owner/user identifiers.",
        "Use composite indexes for high-volume filtered reports, especially `(DeleteStatus, DateColumn)`, `(DeleteStatus, Status_Id)`, and `(Lead_Id, Activity_Date)`.",
        "Avoid over-indexing write-heavy detail tables; review slow query logs before adding broad composite indexes.",
    ])
    page_break(doc)


def api_doc(doc):
    add_h1(doc, "8. API Documentation")
    add_para(doc, "API endpoints are grouped by Express route mounts in `backend/app.js`. Login and selected public print routes are available without JWT; most ERP routes require a Bearer token after the JWT middleware is registered.")
    add_table(doc, ["Method", "Endpoint", "Purpose", "Authentication"], api_groups, [0.65, 2.25, 2.75, 0.9])
    add_h2(doc, "Sample Payloads")
    add_code_block(doc, [
        "POST /Lead/Save_Lead/",
        "{",
        '  "Lead_Id": 0,',
        '  "Lead_Name": "ABC Industries",',
        '  "Phone": "9999999999",',
        '  "Contact_Person": "Operations Head",',
        '  "Email": "contact@example.com",',
        '  "Status_Id": 1,',
        '  "Staff_Id": 12,',
        '  "Is_FollowUp": 1,',
        '  "FollowUp_Date": "2026-06-05",',
        '  "Login_User_Id": 1',
        "}",
        "",
        "200 OK",
        '{ "success": true, "message": "Saved Successfully", "data": { "Key_Id": 101 } }',
    ])
    add_table(doc, ["Status Code", "Meaning", "Expected Client Handling"], [
        ("200", "Successful request.", "Display success or render returned data."),
        ("201", "Created record, where implemented.", "Refresh list/detail state."),
        ("400", "Validation failure or malformed request.", "Show field-level message and keep user input."),
        ("401", "Missing/invalid token.", "Redirect to login and clear session."),
        ("403", "Authenticated but not permitted.", "Show access-denied state."),
        ("404", "Route or record not found.", "Show not-found or refresh listing."),
        ("500", "Database/server error.", "Show safe error message and log details server-side."),
    ], [0.9, 2.2, 3.0])
    page_break(doc)


def permissions_security(doc):
    add_h1(doc, "9. User Roles & Permissions")
    roles = ["Super Admin", "Admin", "Manager", "Sales Executive", "Inventory Manager", "Accounts Team", "Employee"]
    rows = []
    access = {
        "Authentication & User Management": ["Full", "Full", "View", "Own", "Own", "Own", "Own"],
        "Dashboard": ["Full", "Full", "Team", "Sales", "Stock", "Finance", "Own"],
        "Lead Management": ["Full", "Full", "Team", "Create/Edit", "View", "View", "Assigned"],
        "Customer Management": ["Full", "Full", "Team", "Create/Edit", "View", "View/Edit", "View"],
        "Sales Management": ["Full", "Full", "Approve", "Create/Edit", "View", "Invoice/Payment", "View"],
        "Purchase Management": ["Full", "Full", "Approve", "View", "Receive", "Payment", "View"],
        "Inventory Management": ["Full", "Full", "Approve", "View", "Full", "View", "View"],
        "Requirement Management": ["Full", "Full", "Approve", "Create/Edit", "Stock View", "View", "Assigned"],
        "Reports Module": ["Full", "Full", "Team", "Sales", "Stock", "Finance", "Own"],
        "Settings Module": ["Full", "Configure", "View", "No", "Limited", "Limited", "No"],
    }
    for module, vals in access.items():
        rows.append(tuple([module] + vals))
    add_table(doc, ["Module"] + roles, rows, [1.35, 0.75, 0.7, 0.75, 0.9, 0.9, 0.8, 0.65])
    add_h1(doc, "10. Security Documentation")
    add_bullets(doc, [
        "JWT authentication: Express JWT validates HS256 Bearer tokens on protected routes after the public login route.",
        "Password encryption: production deployment should store salted password hashes using bcrypt/argon2 rather than reversible/plain values.",
        "Role-based access: Angular guards and user/menu selection control navigation; backend routes remain token-protected.",
        "Session management: client stores token for active session and clears it on logout/token expiry.",
        "SQL injection prevention: use parameterized MySQL queries and stored procedures; avoid string concatenation for user inputs.",
        "XSS prevention: Angular template binding helps reduce DOM injection; sanitize rich text and never render untrusted HTML directly.",
        "CSRF protection: Bearer-token APIs reduce cookie CSRF exposure; add CSRF tokens if cookie-based auth is introduced.",
        "API security: enforce HTTPS, CORS allow-listing, request-size limits, rate limiting on login, and safe error responses.",
        "Audit logs: preserve login, lead activity, workflow approvals, financial transaction changes, and administrative permission changes.",
    ])
    page_break(doc)


def workflows_deploy_test(doc):
    add_h1(doc, "11. Workflow Diagrams")
    flows = {
        "Login Flow": ["User opens Angular login", "Credentials submitted to Login API", "Backend validates user in MySQL", "JWT signed and returned", "Angular stores token", "Guard permits AdminModule route"],
        "Lead Flow": ["Create lead", "Assign department/status/staff", "Record follow-up", "Log activity", "Schedule meeting or quote", "Convert to requirement/sales opportunity or close as lost"],
        "Sales Flow": ["Prepare quotation", "Confirm quotation", "Create sales order", "Generate invoice/delivery order", "Record receipt/payment reference", "Report sales and outstanding"],
        "Purchase Flow": ["Create supplier/purchase requirement", "Prepare purchase order", "Receive goods/GRN", "Update stock", "Record supplier payment", "Report purchase and VAT"],
        "Inventory Flow": ["Maintain item master", "Receive stock", "Adjust/transfer stock", "Perform stock take", "Reconcile variance", "Publish stock report"],
        "Approval Flow": ["Request created", "Manager reviews", "Approve/reject/return", "Workflow status recorded", "Requester notified", "Transaction proceeds"],
    }
    for name, steps in flows.items():
        add_h2(doc, name)
        add_code_block(doc, [" -> ".join(steps)])
        add_numbered(doc, steps)
    add_h1(doc, "12. Deployment Architecture")
    add_table(doc, ["Environment", "Purpose", "Recommended Configuration"], [
        ("Development", "Local coding and module verification.", "Angular `ng serve`, Node nodemon, local MySQL copy, debug logging enabled."),
        ("Testing/UAT", "Client validation and regression testing.", "Production-like database clone, seeded users/roles, HTTPS test domain, controlled release build."),
        ("Production", "Live business operation.", "Nginx static Angular hosting, Node API behind reverse proxy, MySQL backups, SSL, monitoring, PM2/systemd."),
    ], [1.2, 2.1, 3.0])
    add_bullets(doc, [
        "Angular build: run the configured production build script; serve generated `dist` assets through Nginx.",
        "Node deployment: install backend dependencies, configure environment-specific DB/JWT secrets, run API on port 3504 or approved internal port.",
        "MySQL configuration: separate users per environment, restrict remote access, apply migration scripts, configure buffer/cache settings.",
        "Nginx setup: route `/api` or backend path prefix to Node service, enable gzip/static caching for Angular assets, and redirect HTTP to HTTPS.",
        "SSL/domain mapping: bind client domain to Nginx, install certificate, enforce TLS 1.2+, and monitor expiry.",
    ])
    add_h1(doc, "13. Testing Documentation")
    add_table(doc, ["Test Type", "Purpose", "Sample Coverage"], [
        ("Unit Testing", "Validate functions/components in isolation.", "Angular components/services, Node helper functions, model validation."),
        ("Integration Testing", "Validate API/database interaction.", "Login, lead save/list/delete, sales order save, purchase order save, stock update."),
        ("System Testing", "Validate full ERP workflows.", "Lead-to-quotation, quotation-to-order, PO-to-GRN-to-stock, invoice-to-receipt."),
        ("UAT Testing", "Validate with business users.", "Role-specific acceptance scenarios and report reconciliation."),
        ("Regression Testing", "Prevent existing workflows from breaking.", "Smoke suite across login, dashboard, masters, transactions, reports."),
    ], [1.2, 2.0, 3.2])
    add_table(doc, ["ID", "Scenario", "Expected Result"], [
        ("TC-001", "Login with valid credentials.", "JWT token returned; admin dashboard opens based on role."),
        ("TC-002", "Login with invalid credentials.", "No token returned; user remains on login page with error."),
        ("TC-003", "Create lead with mandatory fields.", "Lead saved and visible in lead list."),
        ("TC-004", "Assign lead follow-up staff/date.", "Latest follow-up appears in lead list and history."),
        ("TC-005", "Create quotation from requirement.", "Quotation header/details saved and printable."),
        ("TC-006", "Receive purchase stock.", "Stock balance increases and purchase/GRN records remain linked."),
        ("TC-007", "Attempt restricted module as Employee.", "Access denied or menu hidden based on permission."),
        ("TC-008", "Export stock report.", "Excel/PDF generated with selected filters."),
    ], [0.75, 3.25, 2.45])
    page_break(doc)


def operations(doc):
    add_h1(doc, "14. Error Handling")
    add_bullets(doc, [
        "Frontend errors: show user-safe messages through toast/dialog components and preserve form state where possible.",
        "Backend errors: route handlers return HTTP 500 with safe JSON for database/internal exceptions; central error handler captures middleware errors.",
        "API errors: standardize response shape with success/message/error fields and status codes.",
        "Database errors: log SQL/procedure errors server-side; return safe message to users without exposing credentials or query internals.",
        "Logging: use Morgan for HTTP logs and structured application logs for critical workflow events, errors, and audit actions.",
    ])
    add_h1(doc, "15. Performance Optimization")
    add_bullets(doc, [
        "Lazy loading: keep authentication and admin modules route-loaded; split larger admin areas if bundle size grows.",
        "Code splitting: production Angular builds should use optimized chunks and gzip compression.",
        "Query optimization: use stored procedures carefully, review slow query logs, and avoid full-table scans in reports.",
        "Caching: cache stable lookup data such as countries, HSN, brands, roles, departments, and designations where appropriate.",
        "Indexing: index join, date, status, search, and soft-delete columns.",
        "API optimization: paginate list endpoints, apply server-side filters, and avoid sending unnecessary large payloads.",
    ])
    add_h1(doc, "16. Backup & Recovery")
    add_table(doc, ["Area", "Strategy", "Frequency / RPO"], [
        ("Database Backup", "Automated MySQL dump or physical backup with encryption and retention.", "Daily full, hourly binlog where feasible."),
        ("File Uploads", "Backup upload/template directories with checksum validation.", "Daily incremental."),
        ("Application Code", "Source control with tagged releases.", "Every release."),
        ("Disaster Recovery", "Restore database, files, backend, frontend build, DNS/SSL, and configuration on standby host.", "RTO defined by SLA."),
    ], [1.3, 3.2, 1.6])
    add_h1(doc, "17. Maintenance Guide")
    add_bullets(doc, [
        "Support process: triage incidents by severity, reproduce issue, capture user role/module/document number, assign owner, and communicate ETA.",
        "Bug fix process: create branch, implement scoped fix, run regression tests, review, deploy to UAT, obtain approval, release to production.",
        "Release process: version code, document migrations, backup production database, deploy backend/frontend, smoke test, monitor logs.",
        "Upgrade process: evaluate Angular/Node/MySQL dependency risks in a staging branch before production adoption.",
    ])
    add_h1(doc, "18. Future Enhancements")
    add_bullets(doc, [
        "Mobile application: formalize the existing Cordova/mobile traces into a maintained Android/iOS application or modern PWA.",
        "AI integration: add lead scoring, sales forecasting, anomaly detection, and document summarization.",
        "Analytics dashboard: introduce richer drill-down BI views and scheduled reports.",
        "Workflow automation: configurable approval matrices, notifications, SLA escalation, and rule-driven task assignment.",
        "Multi-tenant architecture: isolate organizations with tenant identifiers, tenant-aware roles, and database-level safeguards.",
    ])
    add_h1(doc, "19. Conclusion")
    add_para(doc, "NEOTRONIC provides a strong foundation for an enterprise ERP platform by combining business modules, controlled access, transactional persistence, reporting, and extensible route/model organization. With disciplined deployment, testing, security hardening, and maintenance practices, it can support operational control, management visibility, audit readiness, and future digital transformation initiatives.")


def appendices(doc):
    page_break(doc)
    add_h1(doc, "Appendix A - Backend Route Inventory")
    route_names = [name for name, _, _, _ in api_groups]
    route_modules = [
        "Login", "Dashboard", "Lead", "LeadRequirement", "requirementmaster", "requirementdetails", "requirementworkflow", "Price_Response",
        "Client_Accounts", "Company", "User_Details", "User_Role", "User_Type", "User_Menu_Selection", "Menu",
        "Sales_Master", "Sales_Details", "salesquotationmaster", "salesquotationdetails", "salesordermaster", "salesorderdetails",
        "Purchase_Master", "Purchase_Details", "Purchase_Order_Master", "purchaseorderdetails", "Stock", "Stock_Details",
        "Stock_Add_Master", "Stock_Add_Details", "Stock_Take_Master", "Receipt_Voucher", "Payment_Voucher", "Journal_Entry",
        "Account_Group", "Account_Years", "Accounts", "Brand", "Item", "Item_Group", "Model", "HSN", "Department", "Designation",
    ]
    rows = [(m, f"/{m}", "Express router mounted in backend/app.js", "Bearer JWT except login/public exceptions") for m in route_modules]
    add_table(doc, ["Module", "Route Prefix", "Description", "Security"], rows, [1.45, 1.3, 2.5, 1.1])
    add_h1(doc, "Appendix B - Glossary")
    add_table(doc, ["Term", "Definition"], [
        ("ERP", "Enterprise Resource Planning; a system for integrated management of business processes."),
        ("RBAC", "Role-Based Access Control; permissions granted by user role and menu assignment."),
        ("JWT", "JSON Web Token; signed token used by the API to authenticate requests."),
        ("Master Data", "Reusable reference data such as users, customers, items, departments, and accounts."),
        ("Transaction Header", "Primary record for a business document such as invoice, order, purchase, or receipt."),
        ("Transaction Detail", "Line-level records belonging to a transaction header."),
        ("Soft Delete", "Marking a record as deleted through a status flag instead of physically removing it."),
        ("UAT", "User Acceptance Testing conducted by business users before release."),
    ], [1.5, 4.9])


def detailed_appendices(doc):
    page_break(doc)
    add_h1(doc, "Appendix C - Detailed Module Playbooks")
    for idx, (name, summary) in enumerate(modules, start=1):
        add_h2(doc, f"C.{idx} {name}")
        add_para(doc, summary)
        add_table(doc, ["Topic", "Client Handover Detail"], [
            ("Business Owner", "The module should have a nominated process owner who confirms policy, roles, approvals, and report interpretation."),
            ("Primary Data", "Header identifiers, transaction dates, status values, owner/user references, remarks, and related master data must be maintained consistently."),
            ("Validation", "Client-side validation improves usability; backend validation and stored procedure checks protect data integrity."),
            ("Auditability", "Important changes should capture user, timestamp, old value, new value, document reference, and reason/remarks."),
            ("Training Notes", "Users should be trained on create/edit/search/export behavior, approval implications, and how module data affects reports."),
            ("Support Notes", "Support teams should collect screen name, role, document number, browser, request payload where safe, and backend log timestamp."),
        ], [1.4, 5.0])
        add_h3(doc, "Data Entry Checklist")
        add_bullets(doc, [
            "Confirm the user is operating under the correct company, financial year, branch/location context, and role.",
            "Confirm related master data exists before creating transaction records.",
            "Verify numeric fields such as quantity, price, discount, tax, total, and balance use agreed precision.",
            "Confirm dates are stored and displayed in the expected business timezone and format.",
            "Attach or reference supporting documents where the process requires proof or approval evidence.",
        ])
        add_h3(doc, "Control Checklist")
        add_bullets(doc, [
            "Segregate creator, approver, and financial posting responsibilities for high-value transactions.",
            "Review deletion/cancellation permissions monthly.",
            "Reconcile module reports with accounting/stock/customer master data at month end.",
            "Archive exported reports and approval evidence according to client retention policy.",
        ])
        page_break(doc)

    add_h1(doc, "Appendix D - API Catalog by Business Area")
    api_catalog = [
        ("Authentication", ["/Login/Login_Check", "/Login/Login_Checks"]),
        ("CRM and Leads", ["/Lead/Save_Lead", "/Lead/Get_Leads", "/Lead/Get_Dropdowns_Lead", "/Lead/Get_Lead_Activity_Log", "/Lead/Save_Lead_Meeting", "/Lead/Save_Lead_Quote_Tracking"]),
        ("Requirements", ["/LeadRequirement/*", "/requirementmaster/*", "/requirementdetails/*", "/requirementworkflow/*", "/Price_Response/*"]),
        ("Sales", ["/salesquotationmaster/*", "/salesquotationdetails/*", "/salesordermaster/*", "/salesorderdetails/*", "/Sales_Master/*", "/Sales_Details/*", "/deliveryordermaster/*", "/deliveryorderdetails/*"]),
        ("Purchase", ["/Purchase_Master/*", "/Purchase_Details/*", "/Purchase_Order_Master/*", "/purchaseorderdetails/*", "/purchaseordermaster/*", "/purchaseorder_purchasemaster/*", "/Purchase_Return_Master/*"]),
        ("Inventory", ["/Item/*", "/Item_Group/*", "/Brand/*", "/Model/*", "/HSN/*", "/Stock/*", "/Stock_Details/*", "/Stock_Add_Master/*", "/Stock_Take_Master/*"]),
        ("Accounts", ["/Accounts/*", "/Account_Group/*", "/Journal_Entry/*", "/Contra_Entry/*", "/Receipt_Voucher/*", "/Payment_Voucher/*", "/Receipt_Reference/*", "/Payment_Reference/*"]),
        ("Administration", ["/User_Details/*", "/User_Role/*", "/User_Type/*", "/User_Menu_Selection/*", "/Menu/*", "/Company/*", "/General_Settings/*"]),
        ("Reports", ["/Dashboard/*", "/Stock_Report screens", "/Ledger screens", "/OutstandingReport screens", "/Vat_Report screens", "/ProfitAndLossReport screens"]),
    ]
    for name, endpoints in api_catalog:
        add_h2(doc, name)
        add_table(doc, ["Endpoint Pattern", "Method Pattern", "Payload / Response Expectations"], [
            (endpoint, "GET/POST by route implementation", "Requests use JSON body or URL parameters; responses return MySQL row sets, success objects, or safe error JSON.") for endpoint in endpoints
        ], [2.3, 1.3, 2.8])
        add_bullets(doc, [
            "All non-public endpoints must include `Authorization: Bearer <token>`.",
            "List endpoints should support pagination/filtering as data volume grows.",
            "Save endpoints should return a stable key identifier and a clear success/failure message.",
            "Delete endpoints should prefer soft-delete for business records with audit or reporting relevance.",
        ])
        page_break(doc)

    add_h1(doc, "Appendix E - Database Table Detail Cards")
    for idx, (table, purpose, pk, rel) in enumerate(tables, start=1):
        add_h2(doc, f"E.{idx} {table}")
        add_table(doc, ["Attribute", "Documentation"], [
            ("Purpose", purpose),
            ("Primary Key", pk),
            ("Foreign Keys / Related Tables", rel),
            ("Common Columns", "Identifier, name/reference fields, date fields, amount/quantity fields where applicable, status fields, user references, remarks, DeleteStatus."),
            ("Data Type Guidance", "Integer identifiers, VARCHAR for codes/names, TEXT for long remarks, DATE/DATETIME for business timestamps, DECIMAL(18,2) for financial values, TINYINT for flags."),
            ("Indexing Guidance", "Index primary key, join columns, status/date filters, document numbers, and DeleteStatus for active-record searches."),
            ("Retention Guidance", "Retain active and historical records according to client audit, tax, finance, and operational retention policy."),
        ], [1.55, 4.85])
        if idx % 2 == 0:
            page_break(doc)

    add_h1(doc, "Appendix F - Extended Test Case Matrix")
    test_rows = []
    areas = ["Login", "Dashboard", "Lead", "Customer", "Quotation", "Sales Order", "Invoice", "Purchase Order", "GRN", "Stock", "Requirement", "Reports", "Settings", "Permissions"]
    counter = 1
    for area in areas:
        for scenario in ["Create/Open", "Edit/Update", "Search/List", "Permission Check"]:
            test_rows.append((f"ETC-{counter:03d}", area, scenario, "Expected result matches role, validation, persistence, and reporting rules."))
            counter += 1
    for i in range(0, len(test_rows), 12):
        add_table(doc, ["Test ID", "Area", "Scenario", "Expected Result"], test_rows[i:i+12], [0.8, 1.25, 1.55, 2.8])
        page_break(doc)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.text = "NEOTRONIC ERP Project Documentation | Confidential"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(100, 100, 100)


def main():
    doc = Document()
    configure(doc)
    cover(doc)
    intro_sections(doc)
    architecture(doc)
    tech_stack(doc)
    functional_modules(doc)
    database_doc(doc)
    api_doc(doc)
    permissions_security(doc)
    workflows_deploy_test(doc)
    operations(doc)
    appendices(doc)
    detailed_appendices(doc)
    add_footer(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
