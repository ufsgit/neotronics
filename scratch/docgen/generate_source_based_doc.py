import json
import re
from collections import defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
INV = ROOT / "scratch" / "docgen" / "source_inventory.json"
OUT = ROOT / "NEOTRONIC_Source_Based_Project_Documentation.docx"

BLUE = RGBColor(31, 78, 121)
MID_BLUE = RGBColor(46, 116, 181)
LIGHT_FILL = "E8EEF5"
GRAY_FILL = "F2F4F7"


def load():
    return json.loads(INV.read_text(encoding="utf-8"))


def set_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False, color=None, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run("" if text is None else str(text))
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, headers, rows, widths=None, font_size=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    tr_pr = t.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for i, h in enumerate(headers):
        set_fill(t.rows[0].cells[i], LIGHT_FILL)
        cell_text(t.rows[0].cells[i], h, True, BLUE, font_size)
        if widths:
            t.rows[0].cells[i].width = Inches(widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cell_text(cells[i], v, False, None, font_size)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return t


def para(doc, text):
    p = doc.add_paragraph(str(text))
    p.paragraph_format.space_after = Pt(6)
    return p


def bullets(doc, items):
    for x in items:
        if x:
            doc.add_paragraph(str(x), style="List Bullet")


def nums(doc, items):
    for x in items:
        doc.add_paragraph(str(x), style="List Number")


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def code(doc, lines):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.cell(0, 0)
    set_fill(c, "F8FAFC")
    c.text = ""
    p = c.paragraphs[0]
    for i, line in enumerate(lines):
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(8)
        if i != len(lines) - 1:
            r.add_break()
    doc.add_paragraph()


def page(doc):
    doc.add_page_break()


def configure(doc):
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in [("Heading 1", 15, MID_BLUE), ("Heading 2", 12.5, MID_BLUE), ("Heading 3", 10.5, BLUE)]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(8)
        st.paragraph_format.space_after = Pt(4)


def clean(s):
    if not s:
        return ""
    return re.sub(r"\s+", " ", str(s)).strip()


def infer_module(name):
    n = name.lower()
    if any(x in n for x in ["lead"]):
        return "Lead Management"
    if any(x in n for x in ["requirement", "price_request", "price response", "pricerequest", "priceresponse"]):
        return "Requirement and Price Workflow"
    if any(x in n for x in ["quotation", "salesorder", "sales_order", "invoice", "delivery", "packing", "sales_return", "credit_note", "debit_note", "performa"]):
        return "Sales and Fulfilment"
    if any(x in n for x in ["purchase", "grn", "supplier"]):
        return "Purchase Management"
    if any(x in n for x in ["stock", "item", "brand", "model", "hsn", "sale_unit"]):
        return "Inventory and Product Masters"
    if any(x in n for x in ["account", "ledger", "voucher", "journal", "contra", "receipt", "payment", "daybook", "profit", "vat", "outstanding", "statement"]):
        return "Accounts and Finance"
    if any(x in n for x in ["user", "role", "type", "menu"]):
        return "User and Permission Administration"
    if any(x in n for x in ["company", "country", "currency", "department", "designation", "vertical", "terms", "working", "general", "size"]):
        return "System Settings and Masters"
    if "dashboard" in n:
        return "Dashboard"
    if "customer" in n or "client" in n:
        return "Customer and Branch Management"
    return "Other Screens"


def source_summary(doc, inv):
    h1(doc, "Source Analysis Summary")
    para(doc, "This document is generated from the inspected project source code only. It uses Angular routing and component files, dynamic login/menu code, sidebar permissions, frontend services and models, backend Express routes and model files, SQL scripts, and visible templates. Where a field, API, menu item, or relationship is not present in source code, it is not invented.")
    table(doc, ["Source Area", "Count / Finding"], [
        ("Angular root routes", len(inv["app_routes"]["root"])),
        ("Angular auth routes", len(inv["app_routes"]["auth"])),
        ("Angular active admin routes", len(inv["app_routes"]["admin"])),
        ("Parsed admin screens", len(inv["screens"])),
        ("Frontend service files", len(inv["services"])),
        ("Frontend model files", len(inv.get("frontend_models", {}))),
        ("Backend route files", len(inv["backend_routes"])),
        ("Backend endpoints parsed", sum(len(x["endpoints"]) for x in inv["backend_routes"].values())),
        ("Backend model files", len(inv["backend_models"])),
        ("Tables referenced by backend models", len(inv["model_tables"])),
        ("SQL files inspected", len(inv["sql_files"])),
        ("SQL tables found by CREATE/ALTER", len(inv["sql_tables"])),
        ("Dynamic sidebar fallback menu items", ", ".join(inv["sidebar"]["ensured_menu_titles"])),
    ], [2.2, 4.0])
    table(doc, ["File / Mechanism", "Evidence"], [
        ("frontend/src/app/app.routing.ts", "Root redirect to `auth/login`; lazy `AuthModule`; guarded lazy `AdminModule` using `CanAdminGuard`."),
        ("frontend/src/app/modules/admin/admin.routing.ts", "70 active child routes under `AdminComponent`; default child redirects to `/Requirement`; wildcard redirects to `/auth/login`."),
        ("frontend/src/app/modules/auth/login/login.component.ts", "Login builds menu and permissions from `Menu_Service.Get_Menu_Permission(Login_Id)` and stores `Routes_Temp`/`Pointer_Temp`."),
        ("frontend/src/app/components/sidebar/sidebar.component.ts", "Sidebar renders dynamic menu items and exposes `Get_Page_Permission(Menu_Id)` returning View/Save/Edit/Delete permissions."),
        ("backend/app.js", "Express mounts public `/Login`, applies JWT middleware, then mounts protected ERP route prefixes."),
        ("backend/dbconnection.js", "MySQL connection pool targets local database `neotronics_db`."),
    ], [2.5, 3.9])
    page(doc)


def cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NEOTRONIC")
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = BLUE
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Source-Code-Based Project Documentation")
    r.font.size = Pt(16)
    r.font.color.rgb = MID_BLUE
    doc.add_paragraph()
    table(doc, ["Document Attribute", "Value"], [
        ("Document Type", "Combined SRD, TDD, and User Manual"),
        ("Basis", "Existing Angular, Node.js, MySQL, REST API source code in this workspace"),
        ("Prepared Date", date.today().strftime("%d %B %Y")),
        ("Version", "1.0 source-derived"),
        ("Confidentiality", "Confidential client handover and technical documentation"),
    ], [2.0, 4.4])
    para(doc, "No generic ERP modules have been added. All routes, screens, APIs, models, tables, permissions, and workflows documented here are derived from the inspected source tree.")
    page(doc)


def architecture(doc, inv):
    h1(doc, "Technical Architecture")
    code(doc, [
        "Browser / Angular 8 SPA",
        "  app.routing.ts",
        "    / -> auth/login",
        "    /auth -> AuthModule -> LoginComponent",
        "    / -> AdminModule guarded by CanAdminGuard",
        "  AdminComponent shell -> SidebarComponent -> router-outlet -> admin screens",
        "        |",
        "        v",
        "Angular services under frontend/src/app/services",
        "        | HTTP GET/POST",
        "        v",
        "Node.js Express backend/app.js",
        "  /Login public -> JWT token issue",
        "  jwt() middleware -> protected route files",
        "  route file -> model file -> mysql2 pool",
        "        |",
        "        v",
        "MySQL database: neotronics_db, stored procedures and tables",
    ])
    h2(doc, "Angular Architecture")
    bullets(doc, [
        "Angular version is defined in `frontend/package.json` as Angular 8.0.0 with Angular Material/CDK 8.x, RxJS 6.5.2, Google Charts, Bootstrap, and Cordova/Android packaging dependencies.",
        "Application shell routing is split into `AuthModule` and `AdminModule` using lazy route strings.",
        "Admin screens are declared in `admin.module.ts`; active URL mapping is in `admin.routing.ts`.",
        "The sidebar is driven by login menu permissions stored in localStorage rather than a fixed compile-time menu array.",
        "Forms use template-driven `ngModel` and reactive forms in login; Material datepicker, select, autocomplete, paginator, table, dialog, spinner, and checkbox modules are imported.",
    ])
    h2(doc, "Node.js Architecture")
    bullets(doc, [
        "Backend uses Express 4.17.1 with CORS, body-parser, cookie-parser, Morgan, static public hosting, request context, auto response wrapper, JWT middleware, and a central error handler.",
        "Backend `app.js` mounts `/Login` before JWT; all route prefixes mounted after `app.use(jwt())` are protected except explicit JWT exceptions.",
        "There are paired route/model files for most resources, for example `routes/Lead.js` and `models/Lead.js`.",
        "Models use `mysql2` pool queries and many stored procedure calls using `CALL ProcedureName(...)`.",
    ])
    h2(doc, "Authentication and Authorization Flow")
    nums(doc, [
        "User opens `/auth/login`.",
        "Login form validates `userName` and `password`.",
        "Frontend `UserData.login()` calls backend login API.",
        "Backend `/Login` validates credentials through the Login model and signs a JWT using `jsonwebtoken` and `config.secret`.",
        "Login component calls `Menu_Service.Get_Menu_Permission(Login_Id)`.",
        "Menu rows are converted to sidebar route entries containing path, title, Menu_Id, View, Save, Edit, Delete, and Menu_Type.",
        "Routes and pointer table are stored in localStorage as `Routes_Temp` and `Pointer_Temp`.",
        "Sidebar renders only menu entries whose `Menu_Type` is not false.",
        "Screens call `Get_Page_Permission(Menu_Id)` to enable/disable View, Save, Edit, and Delete behavior.",
    ])
    page(doc)


def routes_and_menu(doc, inv):
    h1(doc, "Application Structure, Route Mapping, and Menu Structure")
    h2(doc, "Root and Auth Routes")
    table(doc, ["Area", "Path", "Target", "Guard / Notes"], [
        ("Root", "/", "Redirect to `auth/login`", "pathMatch full"),
        ("Root", "/auth", "Lazy `AuthModule`", "Contains login route"),
        ("Root", "/", "Lazy `AdminModule`", "Protected by `CanAdminGuard`"),
        ("Auth", "/auth/login", "LoginComponent", "Public login screen"),
    ], [1.0, 1.4, 2.4, 1.8])
    h2(doc, "Active Admin Route Mapping")
    rows = [(r["path"], r["component"], r.get("component_import", "")) for r in inv["app_routes"]["admin"]]
    for i in range(0, len(rows), 25):
        table(doc, ["URL", "Component", "Component Import"], rows[i:i+25], [1.6, 2.0, 2.8])
    h2(doc, "Dynamic Login Menu Mapping")
    para(doc, "The following menu rows are generated in `login.component.ts` when backend menu permission rows are returned. Each mapping carries backend fields `VIew_All`, `Menu_Save`, `Menu_Edit`, `Menu_Delete`, and `Menu_Type` into the client route entry.")
    rows = [(m["Menu_Id"], m["title"], m["path"], ", ".join(m["permissions_from_backend"])) for m in inv.get("login_menu_mapping", [])]
    for i in range(0, len(rows), 30):
        table(doc, ["Menu ID", "Menu Title", "Menu Path", "Permission Fields"], rows[i:i+30], [0.7, 1.8, 1.8, 2.1])
    h2(doc, "Site Map")
    code_lines = ["NEOTRONIC", "├── /auth/login - Login"]
    groups = defaultdict(list)
    for r in inv["screens"]:
        groups[infer_module(r["route_path"])].append(r)
    for mod in sorted(groups):
        code_lines.append(f"├── {mod}")
        for r in sorted(groups[mod], key=lambda x: x["path"]):
            code_lines.append(f"│   ├── {r['path']} - {r['component']}")
    code(doc, code_lines[:240])
    page(doc)


def dashboards(doc, inv):
    h1(doc, "Dashboards")
    para(doc, "The source contains one implemented `DashboardComponent` folder and dashboard backend route/model files. `DashboardComponent` is declared in `admin.module.ts` and appears in login menu mapping as Menu_Id 87 with path `/Dashboard`; it is not present as an active child in `admin.routing.ts` in the inspected file.")
    dash = next((s for s in inv["screens"] if s["component"] == "DashboardComponent"), None)
    table(doc, ["Dashboard Source", "Finding"], [
        ("Component files", ", ".join(inv.get("dashboard_files", []))),
        ("Backend route/model", "backend/routes/Dashboard.js and backend/models/Dashboard.js"),
        ("Menu mapping", "Menu_Id 87 -> /Dashboard -> Dashboard"),
        ("Active admin route", "Not found in `admin.routing.ts` active route list."),
    ], [2.0, 4.3])
    para(doc, "Visible dashboard filters and KPIs are taken from `Dashboard.component.html` and `Dashboard.component.ts`.")
    table(doc, ["Dashboard Element", "Source-Derived Detail"], [
        ("Filters", "Look_In_Date checkbox; From date; To date; Branch select; Search button."),
        ("API", "Dashboard_Service.Search_Dashboard_Details(fromDate, toDate, lookInDate, Branch, Login_User)."),
        ("Dropdown source", "Stock_Add_Details_Service.Get_ItemGroup_Load_Data() for branch/client account data."),
        ("KPI card 1", "Purchase Details Count -> `Dashboard_Count`."),
        ("KPI card 2", "Petty Cash Receipt/Payment -> `Dashboard_Count1` / `Dashboard_Count2`."),
        ("KPI card 3", "Payment Voucher Count -> `Dashboard_Count3`."),
        ("KPI card 4", "Waste Management Count -> `Dashboard_Count4`."),
        ("KPI card 5", "Stock Transfer Count -> `Dashboard_Count5`."),
        ("Charts", "Google Chart variables exist for PieChart and ColumnChart, but chart blocks are commented/hidden in current template."),
        ("Permissions", "Uses `Get_Page_Permission(34)` in DashboardComponent, setting Dashboard_Edit/Save/Delete."),
    ], [1.7, 4.7])
    page(doc)


def modules(doc, inv):
    h1(doc, "Module Documentation")
    groups = defaultdict(list)
    for screen in inv["screens"]:
        groups[infer_module(screen["route_path"])].append(screen)
    service_names = set(inv["services"].keys())
    backend_names = set(inv["backend_routes"].keys())
    model_names = set(inv["backend_models"].keys())
    for mod in sorted(groups):
        h2(doc, mod)
        screens = sorted(groups[mod], key=lambda s: s["path"])
        related_names = {s["route_path"].replace("/", "").replace("-", "_") for s in screens}
        apis = []
        dbtables = []
        svcs = []
        for name in sorted(backend_names):
            if any(part.lower() in name.lower() or name.lower() in part.lower() for part in related_names):
                br = inv["backend_routes"][name]
                apis.extend([f"/{name}{e['path']} [{e['method']}]" for e in br["endpoints"][:8]])
        for name in sorted(model_names):
            if any(part.lower() in name.lower() or name.lower() in part.lower() for part in related_names):
                bm = inv["backend_models"][name]
                dbtables.extend(bm.get("tables", []))
        for name in sorted(service_names):
            if any(part.lower() in name.lower() or name.lower() in part.lower() for part in related_names):
                svcs.append(name)
        table(doc, ["Item", "Source-Derived Documentation"], [
            ("Screens", ", ".join(f"{s['path']} ({s['component']})" for s in screens)),
            ("Purpose", f"Business area represented by actual route/component set: {', '.join(s['route_path'] for s in screens[:12])}."),
            ("Frontend services found", ", ".join(svcs[:20]) if svcs else "No directly matched service name found by parser; inspect component service injections for exact calls."),
            ("Backend APIs found", "\n".join(apis[:30]) if apis else "No directly matched backend route name found by parser."),
            ("Database tables referenced", ", ".join(sorted(set(dbtables))[:30]) if dbtables else "No direct model table references found by parser."),
            ("Permissions", "Menu-driven View/Save/Edit/Delete where route is mapped from login menu; screen-level code often calls `Get_Page_Permission(Menu_Id)`."),
        ], [1.55, 4.9])
        page(doc)


def screen_docs(doc, inv):
    h1(doc, "Screen Documentation")
    for idx, s in enumerate(sorted(inv["screens"], key=lambda x: x["path"]), start=1):
        h2(doc, f"{idx}. {s['path']} - {s['component']}")
        templ = s["template"]
        logic = s["component_logic"]
        fields = templ["labels"] + templ["placeholders"] + templ["names"] + templ["ngModels"]
        fields = []
        seen = set()
        for x in templ["labels"] + templ["placeholders"] + templ["names"] + templ["ngModels"]:
            x = clean(x)
            if x and x not in seen:
                seen.add(x)
                fields.append(x)
        table(doc, ["Attribute", "Source-Derived Detail"], [
            ("Route URL", s["path"]),
            ("Component", s["component"]),
            ("Component file", s.get("ts_file", "")),
            ("Template file", s.get("html_file", "")),
            ("Purpose", f"Screen represented by `{s['component']}` and route `{s['path']}`."),
            ("Fields / labels / models", ", ".join(fields[:45]) if fields else "No visible form fields parsed from template."),
            ("Dropdowns", ", ".join(templ["selects"][:30]) if templ["selects"] else "No select/mat-select controls parsed."),
            ("Buttons / links", ", ".join(templ["buttons"][:30]) if templ["buttons"] else "No button text parsed."),
            ("Click actions", ", ".join(templ["click_handlers"][:30]) if templ["click_handlers"] else "No click handlers parsed."),
            ("Component methods", ", ".join(logic["methods"][:35]) if logic["methods"] else "No methods parsed."),
            ("Injected services", ", ".join(logic["services"][:20]) if logic["services"] else "No explicit service injections parsed."),
            ("Service method calls", ", ".join(logic["service_method_calls"][:25]) if logic["service_method_calls"] else "No direct `Service.method()` calls parsed."),
            ("Validations", ", ".join(logic["validations"][:15]) if logic["validations"] else "No explicit validation expressions parsed."),
        ], [1.5, 4.9], font_size=7.8)
        if idx % 2 == 0:
            page(doc)


def database(doc, inv):
    h1(doc, "Database Analysis")
    para(doc, "Database documentation below is limited to tables/columns/procedures visible in backend model queries, SQL files, and frontend model definitions. A full live MySQL schema dump was not available in source, so columns are listed where they are present in SQL CREATE/ALTER scripts or frontend model classes.")
    rows = []
    for t in sorted(set(inv["model_tables"]) | set(inv["sql_tables"].keys())):
        cols = inv["sql_tables"].get(t, [])
        rows.append((t, "Referenced in backend models or SQL scripts", ", ".join(f"{c['name']} {c['type']}" for c in cols[:20]) if cols else "Columns not fully defined in source SQL; see model/frontend field references.", "Used by matching route/model queries or stored procedures."))
    for i in range(0, len(rows), 25):
        table(doc, ["Table", "Purpose Evidence", "Columns Found in Source", "Business Usage"], rows[i:i+25], [1.4, 1.7, 2.1, 1.2], font_size=7.5)
    h2(doc, "Frontend Model Field Definitions")
    model_rows = []
    for name, m in sorted(inv.get("frontend_models", {}).items()):
        model_rows.append((name, m["file"], ", ".join(f"{f['name']}: {f['type']}" for f in m["fields"][:30]) or "No fields parsed"))
    for i in range(0, len(model_rows), 20):
        table(doc, ["Model", "File", "Fields"], model_rows[i:i+20], [1.3, 2.0, 3.1], font_size=7.4)
    page(doc)


def apis(doc, inv):
    h1(doc, "API Documentation")
    para(doc, "API endpoint documentation is parsed from Express route files. Request and response payloads are documented as source-derived where route/model code exposes body, params, or stored procedure usage; detailed schemas should be confirmed per endpoint during API contract hardening.")
    for name, br in sorted(inv["backend_routes"].items()):
        h2(doc, f"/{name}")
        model = inv["backend_models"].get(br["model"], {})
        table(doc, ["Attribute", "Detail"], [
            ("Route file", br["file"]),
            ("Model", br.get("model", "")),
            ("Model file", model.get("file", "")),
            ("Model methods", ", ".join(model.get("methods", [])[:50])),
            ("Stored procedures", ", ".join(model.get("procedures", [])[:50])),
            ("Tables referenced", ", ".join(model.get("tables", [])[:50])),
            ("Error handling pattern", "Route handlers commonly return JSON errors or HTTP 500 for database/internal exceptions; global error handler also registered in backend/app.js."),
        ], [1.4, 5.0], font_size=7.8)
        rows = []
        for e in br["endpoints"]:
            request = "Uses req.body for POST routes; req.params for `:id` path segments; req.query where route code references query."
            response = "JSON response from model callback, stored procedure row sets, or success/error object."
            rows.append((e["method"], f"/{name}{e['path']}", request, response))
        for i in range(0, len(rows), 18):
            table(doc, ["Method", "Endpoint", "Request Source", "Response Source"], rows[i:i+18], [0.65, 2.0, 1.85, 1.9], font_size=7.2)
        page(doc)


def workflows(doc):
    h1(doc, "Workflow Documentation")
    flows = {
        "User Login": ["Open /auth/login", "Enter userName/password", "Validate Angular form", "Call login API", "Store token/Login_User", "Fetch menu permissions", "Build ROUTES and Pointer_Table", "Navigate to Lead"],
        "Lead Management": ["Open /Lead", "Load dropdowns/stages", "Create or edit lead", "Assign staff/status/follow-up", "Save through Lead API", "Log activity/follow-up/meeting/quote tracking", "Review lead list/history"],
        "Requirement Management": ["Open /Requirement or /LeadRequirement", "Capture requirement header/details", "Use requirementworkflow for status progression", "Use PriceRequest/PriceResponse where pricing is needed", "Convert to quotation/order flow where applicable"],
        "Customer Management": ["Open /Customer or /Client_Accounts", "Maintain branch/client account data", "Use records in dashboard branch filters, sales, receipts, ledger, and statements"],
        "Quotation Process": ["Open /Quotation", "Select customer/requirement/items", "Save quotation master/details", "Use /Quotation_Confirmation where required", "Proceed to SalesOrder/Delivery/Invoice screens"],
        "Sales Process": ["Open SalesOrder/Invoice/Delivery_Order", "Create order/invoice/delivery records", "Use receipt/credit/debit note screens for financial follow-up", "Report through sales/accounting reports"],
        "Purchase Process": ["Open Purchase_order/Purchase_Master/GRN", "Create purchase order", "Receive goods through GRN", "Record purchase master/details", "Use purchase return where needed"],
        "Inventory Process": ["Maintain Item/Item_Group/Brand/Model/HSN/Sale_Unit", "Use AddStock/Stock_Adjust/Stock_Take", "Review Stock_Reports", "Use item and stock references in sales/purchase/requirements"],
        "Approval Process": ["Permissions and menu visibility arrive from backend menu permission rows", "Screen Save/Edit/Delete rights are resolved with Get_Page_Permission", "Requirement workflow has dedicated backend route/model", "Quotation confirmation has a dedicated screen route/menu mapping"],
        "Reporting Process": ["Open report screen from dynamic menu", "Apply filters/date/branch/account/item", "Fetch report data through related service/API", "View table/cards", "Export where screen provides export buttons/libraries"],
    }
    for name, steps in flows.items():
        h2(doc, name)
        code(doc, [f"{i+1}. {step}" for i, step in enumerate(steps)])
        nums(doc, steps)
    page(doc)


def roles(doc, inv):
    h1(doc, "User Roles and Permissions")
    para(doc, "No fixed named business-role list is hard-coded in the inspected routing layer. The source implements permission control through `User_Role`, `User_Type`, `User_Menu_Selection`, `Menu`, login menu permission rows, and sidebar pointer lookup.")
    table(doc, ["Permission Source", "Finding"], [
        ("Backend resources", "Routes/models exist for User_Role, User_Type, User_Details, User_Menu_Selection, and Menu."),
        ("Login fields used", "Menu rows expose `VIew_All`, `Menu_Save`, `Menu_Edit`, `Menu_Delete`, and `Menu_Type`."),
        ("Client storage", "`Routes_Temp` stores menu entries; `Pointer_Temp` maps Menu_Id to route array index."),
        ("Screen check", "`Get_Page_Permission(Menu_Id)` returns View, Save, Edit, Delete values."),
        ("Route guard", "`CanAdminGuard` allows admin module only when `UserData.isLoggedIn()` is true."),
    ], [1.6, 4.8])
    rows = [(m["Menu_Id"], m["title"], m["path"], "View/Save/Edit/Delete/Menu_Type from backend") for m in inv.get("login_menu_mapping", [])]
    for i in range(0, len(rows), 30):
        table(doc, ["Menu ID", "Screen/Menu", "Path", "Permission Model"], rows[i:i+30], [0.7, 1.8, 1.8, 2.1], font_size=7.5)
    page(doc)


def handover(doc, inv):
    h1(doc, "Client Handover Summary")
    para(doc, "NEOTRONIC is an Angular 8 and Node.js/MySQL ERP-style web application with dynamic menu-based authorization, a broad set of business screens, and route/model API resources for sales, purchase, inventory, accounts, lead, requirement, report, and settings operations.")
    bullets(doc, [
        "Client-visible navigation is determined by backend menu permissions at login.",
        "The application shell supports guarded authenticated administration and a public login module.",
        "The backend API surface is large: 87 route files and 820 parsed route endpoints.",
        "Screens use Angular components and HTML templates under `frontend/src/app/modules/admin`.",
        "Database access is concentrated in backend model files and stored procedures.",
    ])
    h2(doc, "Future Enhancements Evident from Source Review")
    bullets(doc, [
        "Add Dashboard to active `AdminRoutes` if the menu path `/Dashboard` is intended to be directly navigable.",
        "Normalize menu paths where login mapping uses names not present in active admin routes, such as `/Sales Order` versus `/SalesOrder`.",
        "Generate an OpenAPI specification from route/model contracts and standardize request/response schemas.",
        "Add complete SQL schema dump or migrations for all referenced tables, not only incremental SQL files.",
        "Centralize permission constants/Menu_Id mappings to reduce drift between login menu mapping and routes.",
    ])


def footer(doc):
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.text = "NEOTRONIC Source-Based SRD/TDD/User Manual | Confidential"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(100, 100, 100)


def main():
    inv = load()
    doc = Document()
    configure(doc)
    cover(doc)
    source_summary(doc, inv)
    architecture(doc, inv)
    routes_and_menu(doc, inv)
    dashboards(doc, inv)
    modules(doc, inv)
    screen_docs(doc, inv)
    workflows(doc)
    database(doc, inv)
    apis(doc, inv)
    roles(doc, inv)
    handover(doc, inv)
    footer(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
